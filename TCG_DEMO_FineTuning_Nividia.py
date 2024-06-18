######### 1. Load requiered libraries  #######################
import sys
import os
import re
import requests
import openpyxl
import docx2txt
import ast
from datetime import datetime
import docx
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from prettytable import PrettyTable
from PyPDF2 import PdfReader
import json
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

from PyQt5 import QtWidgets
from PyQt5.QtCore import Qt, QSize, QThread, QTimer
from PyQt5.QtGui import QPixmap, QFont, QColor, QPalette, QPixmap, QBrush
from PyQt5.QtWidgets import (QApplication, QMainWindow, QProgressDialog,
    QVBoxLayout, QHBoxLayout, QWidget, QGroupBox, QLabel, QPushButton,
    QFrame, QHBoxLayout, QWidget, QFileDialog, QMessageBox, QDialog, QTextEdit,
    QSpinBox, QDialogButtonBox, QComboBox, QProgressBar, QSpacerItem, QSizePolicy,
    QGridLayout, QLineEdit)

##########################################################################
######### UI Temp Classes #######################
## CustomButton - Creates a QPushButton with the desired styling.

class StyledButton(QPushButton):
    def __init__(self, text,width=200,height=35,background_color="#29465B"):
        self.background_color= background_color
        self.width = width
        self.height = height
        super().__init__(text)

        self.setFixedSize(self.width, self.height)
        self.setStyleSheet(f"""
        background-color: {self.background_color}; 
        color: white; font-weight: bold; border-radius: 5px;
        """)

## CustomLogoLabel - Creates a QLabel for logo images.
class StyledBox(QGroupBox):
    def __init__(self, width, height):
        super().__init__()
        self.setFixedSize(width, height)
        self.setStyleSheet("""
            QGroupBox {
                background-color: #488AC7;
                border-radius: 5px;
                color: white;
             
            }
        """)

### Custom LayoutBox 

class ButtonDialogStyle:
    def __init__(self, background_color="#488AC7", border_color="white", border_width="3px",
                 border_radius="5px", padding="10px 25px", button_background_color="white",
                 button_border_color="white", button_border_radius="5px", button_padding="5px 15px",
                 button_hover_color="lightgrey", button_pressed_color="lightgrey",min_height="1.3em"):
        
        self.background_color = background_color
        self.border_color = border_color
        self.border_width = border_width
        self.border_radius = border_radius
        self.padding = padding
        self.button_background_color = button_background_color
        self.button_border_color = button_border_color
        self.button_border_radius = button_border_radius
        self.button_padding = button_padding
        self.button_hover_color = button_hover_color
        self.button_pressed_color = button_pressed_color
        self.min_height = min_height

    def set_style(self, dialog):
        dialog.setStyleSheet(f"""
            QDialog {{
                border-radius: {self.border_radius};
                background-color: {self.background_color};
                border: {self.border_width} solid {self.border_color};
                min-height: {self.min_height};
                padding: {self.padding};
            }}
            QPushButton {{
                background-color: {self.button_background_color};
                border: 2px solid {self.button_border_color};
                border-radius: {self.button_border_radius};
                min-height: {self.min_height};
                padding: {self.button_padding};
            }}
            QPushButton:hover {{
                background-color: {self.button_hover_color};
            }}
            QPushButton:pressed {{
                background-color: {self.button_pressed_color};
            }}
        """)

class ProgressDialogStyle:
    def __init__(self, background_color="#488AC7", border_color="white", border_width="3px",
                 border_radius="5px", padding="10px 25px", label_color="white", label_font_weight="bold",
                 progress_background_color="white", progress_text_align="right", progress_color="black",
                 progress_chunk_background_color="#46C646", progress_chunk_font_weight="bold"):
        self.background_color = background_color
        self.border_color = border_color
        self.border_width = border_width
        self.border_radius = border_radius
        self.padding = padding
        self.label_color = label_color
        self.label_font_weight = label_font_weight
        self.progress_background_color = progress_background_color
        self.progress_text_align = progress_text_align
        self.progress_color = progress_color
        self.progress_chunk_background_color = progress_chunk_background_color
        self.progress_chunk_font_weight = progress_chunk_font_weight

    def set_style(self, dialog):
        dialog.setStyleSheet(f"""
            QProgressDialog {{
                border-radius: {self.border_radius};
                background-color: {self.background_color};
                border: {self.border_width} solid {self.border_color};
                min-height: 1.3em;
                padding: {self.padding};                 
            }}
            QProgressDialog QLabel {{
                color: {self.label_color};
                font-weight: {self.label_font_weight};
            }}
            QProgressDialog QProgressBar {{
                background-color: {self.progress_background_color};
                text-align: {self.progress_text_align};
                color: {self.progress_color};
            }}
            QProgressDialog QProgressBar::chunk {{
                background-color: {self.progress_chunk_background_color};
                font-weight: {self.progress_chunk_font_weight};
            }}
        """)

class MainWindow(QMainWindow):
    # static variables:

    text = None
    # req_temp = None
    # usr_story = None
    # generic_doc = None
    req_response_data = None
    generate_test_response_data = None
    gen_tst_case = None
    gen_tst_req = None
    summary_response = None
    req_txt = None
    test_cases = None
    #upload_to_opkey = None
     
    ### All UI Buttons 
    upload_button = False
    summary_button = False
    generate_test_requirements_button = False
    generate_test_cases_button = False
    generate_automation_scripts_button = False
    user_input_button = False
    upload_to_opkey_button = False
    export_file_button = False
    requirements_import_button = False
    clear_button = False
    exit_button = False

    ## User Input Option Button (Default Value)
    delivery_team_user_input = '...'
    generate_test_type_user_input = '...' ## by defaulf wil be empty string
    priority_user_input = "..."
    milestone_user_input = "..."

    get_link_1 = "https://demo.labs.opkeyone.com/api/OpkeyAuth/GetListOfAssignedProject"
    post_link_1 = "https://demo.labs.opkeyone.com/api/OpkeyAuth/Login_API_KEY?username=tuhinmitra@kpmg.com&loginsource=TCG"
    post_link_2 = "https://demo.labs.opkeyone.com/api/OpkeyAuth/SelectProject?projectId=9e188d9d-29eb-4dc5-af85-d39638f9e3df"
    post_link_3 = "https://demo.labs.opkeyone.com/Automation/ExecuteOracleTestCreationQuery"
    api_key = "5SOEH6MWEPX6NZ4IX0"
    
    # Init function used to design UI
    def __init__(self):
        
        super().__init__()
        # Get the current working directory
        current_path = os.getcwd()

        # Set the title for the main window
        self.setWindowTitle("Developer: KPMG Team")

        # Maximize the window size on launch
        self.setWindowState(Qt.WindowMaximized)

        # Create a QPalette object to set the background image
        palette = QPalette()
        bg_image_path = os.path.join(current_path, "UI_Temp", "BG-1.jpg")
        bg_image = QPixmap(bg_image_path)
        palette.setBrush(QPalette.Background, QBrush(bg_image))

        # Apply the customized palette to the main window
        self.setPalette(palette)

        # Create a QVBoxLayout for the main layout
        layout = QVBoxLayout()

        # Create a top box group and a horizontal layout for the box
        main_top_box_group = QGroupBox()
        # main_top_box_group.setFixedSize(1900, 110)
        main_top_box_group.setStyleSheet("QGroupBox {background-color: transparent; border: transparent;}")
        main_top_box_layout = QHBoxLayout(main_top_box_group)

        top_box_group = StyledBox(1900, 100)
        top_box_layout = QHBoxLayout(top_box_group)
        top_box_layout.setAlignment(Qt.AlignCenter) ### remove if KMPG logo needs to be left
        

        # Add KPMG logo QLabel using CustomLogoLabel
        kpmg_logo_label = QLabel()
        kpmg_logo_path = os.path.join(current_path, "UI_Temp", "kpmg-logo-2.png")
        kpmg_logo_pixmap = QPixmap(kpmg_logo_path)
        kpmg_logo_label.setPixmap(kpmg_logo_pixmap.scaledToHeight(int(top_box_group.height() * 1.5), Qt.SmoothTransformation))
        top_box_layout.addWidget(kpmg_logo_label, alignment=Qt.AlignLeft)
        top_box_layout.addSpacing(620) 
        # Add the HP logo QLabel using CustomLogoLabel
        # hp_logo_label = QLabel()
        # hp_logo_path = os.path.join(current_path, "UI_Temp", "hp-logo-4.png")
        # hp_logo_pixmap = QPixmap(hp_logo_path)
        # hp_logo_label.setPixmap(hp_logo_pixmap.scaledToHeight(top_box_group.height() * 0.7, Qt.SmoothTransformation))
        # top_box_layout.addWidget(hp_logo_label, alignment=Qt.AlignLeft|Qt.AlignCenter)

        # Add a QLabel widget for the heading and add it to the layout
        heading = QLabel("KPMG Test Generator")
        heading.setFont(QFont("KPMG Bold", 30, QFont.Bold))
        heading.setStyleSheet("color: white;")
        top_box_layout.addWidget(heading, alignment= Qt.AlignLeft|Qt.AlignCenter)
        top_box_layout.addSpacing(800)  # Adding a gap of 20 pixels between buttons
        # Add the top box to the main layout with center alignment
        # layout.addWidget(top_box_group, alignment=Qt.AlignTop)
        # main_top_box_layout.addWidget(top_box_group, alignment=Qt.AlignTop)  #-------

        # Add the disclaimer label to the layout
        disclaimer_label = QLabel()
        disclaimer_path = os.path.join(current_path, "UI_Temp", "disclaimer.png")
        disclaimer_pixmap = QPixmap(disclaimer_path)
        disclaimer_label.setPixmap(disclaimer_pixmap.scaledToHeight(int(top_box_group.height() * .5), Qt.SmoothTransformation))
        disclaimer_label.setToolTip("""        Please note that the Test Case Generator Tool is a software tool designed for generating test cases quickly and efficiently from raw requirements. 
        It is intended for use by software developers and testers who have experience in programming and application development. 
        While the tool is designed to simplify and streamline the test case generation process, 
        it is still necessary for users to have a good understanding of their application's requirements and functionality.
        The generated test cases should be reviewed and validated by an experienced tester before being executed on any application or software system. 
        The tool assumes no responsibility for errors or issues that may arise from the use of generated test cases or automation scripts.
                                    
        The Test Case Generator Tool can help to save lot of time and reduce human errors in the software testing process, 
        making it an essential tool for software developers and testers
        """)
        # Set font weight to bold
        disclaimer_font = QFont()
        disclaimer_font.setBold(True)
        disclaimer_label.setFont(disclaimer_font)

        # Set border style and width
        disclaimer_label.setStyleSheet("color: black; font: bold")
        # layout.addWidget(disclaimer_label, alignment= Qt.AlignTop | Qt.AlignRight )
        # main_top_box_layout.addWidget(disclaimer_label, alignment= Qt.AlignBottom | Qt.AlignRight ) #-----
        top_box_layout.addWidget(disclaimer_label, alignment= Qt.AlignBottom | Qt.AlignRight )

        #Add the maintop box in layout
        # layout.addWidget(main_top_box_group, alignment=Qt.AlignTop) #---------
        layout.addWidget(top_box_group, alignment=Qt.AlignTop)
        
        # # Create a QLabel widget for the subheading and add it to the layout
        # sub_heading = QLabel("A US Digital LightHouse GenAI Innovation")
        # sub_heading.setFont(QFont("KPMG Bold", 15, QFont.Bold))
        # sub_heading.setStyleSheet("color: black;")
        # # sub_heading.setAlignment(Qt.AlignCenter)
        # layout.addWidget(sub_heading, alignment= Qt.AlignCenter | Qt.AlignTop )
        # # box_layout.addWidget(sub_heading, alignment=Qt.AlignCenter)

        hbox = QHBoxLayout()
        MainWindow.upload_button = StyledButton(text="Upload Requirement")
        hbox.addWidget(MainWindow.upload_button)
        MainWindow.upload_button.clicked.connect(self.upload_document)

        MainWindow.summary_button = StyledButton(text="Provide Summary")
        hbox.addWidget(MainWindow.summary_button)
        MainWindow.summary_button.clicked.connect(self.provide_summary)

        MainWindow.generate_test_requirements_button = StyledButton(text="Structure Test Requirements",width=250)
        hbox.addWidget(MainWindow.generate_test_requirements_button)
        MainWindow.generate_test_requirements_button.clicked.connect(self.generate_test_requirements)

        MainWindow.generate_test_cases_button = StyledButton(text="Generate Test Cases",width=200)
        hbox.addWidget(MainWindow.generate_test_cases_button)
        # As intial diabled the button
        MainWindow.generate_test_cases_button.setEnabled(False)
        MainWindow.generate_test_cases_button.setStyleSheet("background-color: #B6B6B4; color: white; font-weight: bold; border-radius: 5px;")
        MainWindow.generate_test_cases_button.clicked.connect(self.test_type)

        # MainWindow.generate_automation_scripts_button = StyledButton(text="Create Automation Scripts",width=230)
        # hbox.addWidget(MainWindow.generate_automation_scripts_button)
        # # As intial diabled the button
        # MainWindow.generate_automation_scripts_button.setEnabled(False)
        # MainWindow.generate_automation_scripts_button.setStyleSheet("background-color: #B6B6B4; color: white; font-weight: bold; border-radius: 5px;")
        # MainWindow.generate_automation_scripts_button.clicked.connect(self.create_automation_scripts)

        MainWindow.user_input_button = StyledButton(text="User Inputs")
        hbox.addWidget(MainWindow.user_input_button)
        MainWindow.user_input_button.clicked.connect(self.user_input)

        MainWindow.export_file_button = StyledButton(text="Export",width=180)
        hbox.addWidget(MainWindow.export_file_button)
        MainWindow.export_file_button.clicked.connect(self.export_button)

        MainWindow.upload_to_opkey_button = StyledButton(text="Upload To Opkey",background_color="#B6B6B4")
        hbox.addWidget(MainWindow.upload_to_opkey_button)
        MainWindow.upload_to_opkey_button.setEnabled(False)
        MainWindow.upload_to_opkey_button.clicked.connect(self.upload_to_opkey_function)

        ### Add Layout
        layout.addLayout(hbox)
        layout.setAlignment(hbox, Qt.AlignCenter)

        self.display_pane = QVBoxLayout()
        self.display_pane.setSpacing(10)
        self.display_pane.setAlignment(Qt.AlignTop)

        self.summary_label = QTextEdit()
        self.summary_label.setReadOnly(True)
        self.summary_label.setStyleSheet("color: black; background-color: white;")
        self.display_pane.addWidget(self.summary_label)

        display_frame = QFrame()
        display_frame.setLineWidth(3)
        display_frame.setFrameShape(QFrame.Box)
        display_frame.setFrameShadow(QFrame.Sunken)
        display_frame.setFixedSize(1500, 600)
        display_frame.setLayout(self.display_pane)
        layout.addWidget(display_frame, alignment=Qt.AlignCenter)

        ## 
        dbox = QHBoxLayout()
        dbox.addSpacing(100)  # Adding a gap of 20 pixels between buttons

        MainWindow.requirements_import_button = StyledButton(text="Requirements Import",width=230,background_color="#B6B6B4")
        dbox.addWidget(MainWindow.requirements_import_button)

        MainWindow.generate_automation_scripts_button = StyledButton(text="Create Automation Scripts",width=230)
        dbox.addWidget(MainWindow.generate_automation_scripts_button)
        # As intial diabled the button
        MainWindow.generate_automation_scripts_button.setEnabled(False)
        MainWindow.generate_automation_scripts_button.setStyleSheet("background-color: #B6B6B4; color: white; font-weight: bold; border-radius: 5px;")
        MainWindow.generate_automation_scripts_button.clicked.connect(self.create_automation_scripts)

        
        dbox.addSpacing(2500)  # Adding a gap of 20 pixels between buttons

        MainWindow.clear_button = StyledButton(text="New Scenario",width=200)
        dbox.addWidget(MainWindow.clear_button)
        MainWindow.clear_button.clicked.connect(self.clear_display)

        exit_button = StyledButton(text="Exit",width=150)
        dbox.addWidget(exit_button)
        exit_button.clicked.connect(self.exit_application)


        dbox.addSpacing(100)  # Adding a gap of 20 pixels between buttons

        layout.addLayout(dbox)
        layout.setAlignment(dbox, Qt.AlignLeft)


        # footer_text = QLabel("This software is the exclusive property of KPMG Lighthouse [Group Name], and unauthorized copying or illicit usage is strictly prohibited. Any violation of these terms may result in legal action.")
        footer_text = QLabel("For Internal Use Only")
        
        footer_text.setAlignment(Qt.AlignCenter)
        footer_text.setFont(QFont("Arial", 9))
        footer_text.setStyleSheet("color: black;")
        layout.addWidget(footer_text)

        central_widget = QWidget()
        central_widget.setLayout(layout)
        self.setCentralWidget(central_widget)

    # method defining T&C popup
    def show_popup(self):
        MainWindow.gen_tst_case = False
        popup_dialog = QDialog(self)
        popup_dialog.setWindowTitle("TCG 1.0")
        popup_dialog.setStyleSheet("background-color: black;")

        # Set fixed size for the QDialog
        popup_dialog.setFixedSize(900, 450)

        vbox = QVBoxLayout()
        popup_dialog.setLayout(vbox)

        heading_label = QLabel("Please read carefully")
        heading_label.setFont(QFont("KPMG Bold", 30))
        heading_label.setStyleSheet("color: red;")
        heading_label.setAlignment(Qt.AlignCenter)
        vbox.addWidget(heading_label)
        message_text = '''\
            This tool, a US Lighthouse GenAI Innovation, assists in generating test cases from requirements and subsequently transforming them into automation scripts. Before using this tool, it's crucial to thoroughly comprehend the requirements. After generating the test cases, a manual review must be conducted. Ensure that you input the ideal number of test cases to achieve a 100% requirement coverage. Please note that the Large Language Model's response could vary due to its nature. Please use your skepticism and professional judgement when using results from this tool in any deliverable.
            If you disagree with any of these terms, please exit the program. For any inquiries, feel free to contact the <Group@xyz.com>.
            '''
        message_label = QLabel(message_text)
        message_label.setFont(QFont("Arial", 9))
        message_label.setStyleSheet("color: red;")
        message_label.setAlignment(Qt.AlignCenter)
        message_label.setWordWrap(True)
        vbox.addWidget(message_label)

        hbox = QHBoxLayout()
        hbox.addStretch()

        accept_button = QPushButton("Accept")
        accept_button.setStyleSheet("background-color: lightblue;")
        hbox.addWidget(accept_button)
        accept_button.clicked.connect(popup_dialog.accept)

        reject_button = QPushButton("Reject")
        reject_button.setStyleSheet("background-color: lightblue;")
        hbox.addWidget(reject_button)
        reject_button.clicked.connect(popup_dialog.reject)

        hbox.addStretch()
        vbox.addLayout(hbox)

        result = popup_dialog.exec_()

        if result == QDialog.Accepted:
            return True
        else:
            return False
    # Method defining document upload button functionality

    def upload_document(self):
        dialog = QDialog(self)
        dialog.resize(600, 150)
        dialog.setWindowTitle("Select the type of Document")

        vbox = QVBoxLayout()
        dialog.setLayout(vbox)

        dialog_style =ButtonDialogStyle()
        dialog_style.set_style(dialog)

        doc_combo = QComboBox()
        doc_combo.setMinimumWidth(200)
        doc_combo.setStyleSheet("background-color: white; border-radius: 5px; min-height: 1.3em; padding: 3px 10px")
        doc_combo.move(60, 30)
        doc_combo.addItem("(default) Select the type of document...")
        doc_combo.addItem("PDF")
        doc_combo.addItem("Word")
        doc_combo.addItem("Text")
        doc_combo.addItem("Excel")
        # doc_combo.addItem("PPT")
        vbox.addWidget(doc_combo)

        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        vbox.addWidget(button_box)

        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)

        result = dialog.exec_()

        ######################

        if result == QDialog.Accepted:
            doc_type = doc_combo.currentText()

            # if doc_type!="(default) Select the type of document...":
            options = QFileDialog.Options()
            options |= QFileDialog.ReadOnly
            if doc_type=="Word":
                file_name, _ = QFileDialog.getOpenFileName(self, "Open Document", "", "Word Documents (*.docx; *.doc)", options=options) 
            elif doc_type=="PDF":
                file_name, _ = QFileDialog.getOpenFileName(self, "Open Document", "", "PDF Documents (*.pdf;)", options=options) 
            elif doc_type=="Text":
                file_name, _ = QFileDialog.getOpenFileName(self, "Open Document", "", "Text Documents (*.txt;)", options=options) 
            elif doc_type=="Excel":
                file_name, _ = QFileDialog.getOpenFileName(self, "Open Document", "", "Excel/CSV Documents (*.xlsx;*.csv;)", options=options) 
            elif doc_type=="PPT":
                file_name, _ = QFileDialog.getOpenFileName(self, "Open Document", "", "PPT Documents (*.pptx;)", options=options) 
            

                

            if file_name: 
                    MainWindow.text = self.extract_text(file_name)
                    if MainWindow.text:
                        progress_dialog = QProgressDialog("Uploading document...", None, 0, 100 , self)
                        progress_dialog.resize(400,100)
                        progress_style = ProgressDialogStyle()
                        progress_style.set_style(progress_dialog)
                        progress_dialog.setWindowTitle("In progress")
                        progress_dialog.setWindowModality(Qt.ApplicationModal)
                        progress_dialog.setAutoClose(False)
                        progress_dialog.setAutoReset(False)
                        progress_dialog.setMinimumDuration(0)
                        progress_dialog.forceShow()  

                        for progress in range(0, 101, 3):
                            progress_dialog.setValue(progress)
                            QApplication.processEvents()
                            progress_dialog.setLabelText(f"Please wait while the document is uploading...")
                            QApplication.instance().processEvents()
                            QThread.msleep(60)
                        # self.send_text_to_gpt(MainWindow.text)
                        progress_dialog.close()
                        QMessageBox.information(self, "Success", "Document successfully uploaded!")
                        MainWindow.upload_button.setEnabled(False)
                        MainWindow.upload_button.setStyleSheet("background-color: green; color: white; font-weight: bold; border-radius: 5px;")
                    else:
                        QMessageBox.critical(self, "Error", "Error extracting text from the document!")
            else:
                QMessageBox.warning(self, "Warning", "Please select the type of document")
    # # Method defining how uploaded content is extracted

    def extract_text(self, file_name):
        _, file_extension = os.path.splitext(file_name)
        if file_extension == '.pdf':
            with open(file_name, 'rb') as file:
                reader = PdfReader(file)
                ext_text = '\n'.join(page.extract_text() for page in reader.pages)
        elif file_extension == '.docx':
            ext_text = docx2txt.process(file_name)
        elif file_extension == '.txt':
            with open(file_name, "r") as file:
                ext_text = file.read()

                # print("--------------------------")
                # print(ext_text)

        elif file_extension == '.xlsx':
            df = pd.read_excel(file_name)
            ext_text = df.to_string(index=False, header=True)  # Convert DataFrame content to a formatted
        else:
            ext_text = None
        return ext_text

    def send_text_to_gpt(self, text):
        self.api_key = "73c315ba0bc44e6abffd10b189dc3199"
        self.url = "https://gpt-llm-4.openai.azure.com/openai/deployments/gpt-4/chat/completions?api-version=2023-03-15-preview"
        #self.api_key = "d4595b37d235460395fa78721602c87b"
        #self.url = "https://oai-kgsgpt-testapp.openai.azure.com/openai/deployments/gpt-35-turbo-16k/chat/completions?api-version=2023-03-15-preview"  

        conversation_history_with_prompt = f"""You are an expert in analyzing functional requirements. 
        Understand the extracted text. 
        Provide summary in not less than 500 words. 
        Summary should be crisp. 
        Display numbers if possible from the extracted text. 
        Dont create numbers on your own: {text}"""

        headers = {"api-key": self.api_key}
        query_data = {
            "messages": [
                {"role": "system", "content": conversation_history_with_prompt}
            ],
            "temperature": 0,
            "top_p": 1,
            "frequency_penalty": 0,
            "presence_penalty": 0,
            "max_tokens": 4096,
            "stop": None,
        }

        MainWindow.summary_response = requests.post(self.url, headers=headers, json=query_data, verify=False)
        print('Summary Response:',MainWindow.summary_response)
        print(str(MainWindow.summary_response))
        summary_response_data = MainWindow.summary_response.json()
        self.store_gpt_response_data(summary_response_data)
        
    '''async def send_chunk_to_gpt(url, headers, query_data):
        response = await requests.post(url, headers=headers, json=query_data)
        response_data = response.json()
        #self.store_gpt_response_data(response_data)
        return response_data'''
    
    '''async def send_text_to_gpt(text, batch_size=5):
            api_key = "d4595b37d235460395fa78721602c87b"
            url = "https://gpt-llm-4.openai.azure.com/openai/deployments/gpt-4/chat/completions?api-version=2023-03-15-preview"

            headers = {"api-key": api_key}
            conversation_history_with_prompt = f"You are an expert in analyzing functional requirements. Understand the extracted text. Provide summary in not less than 500 words. Summary should be crisp. Display numbers if possible from the extracted text. Dont create numbers on your own: {text}"
            query_data = {"messages": [{"role": "system", "content": conversation_history_with_prompt}],
                        "temperature": 0, "top_p": 1, "frequency_penalty": 0, "presence_penalty": 0,
                        "max_tokens": 4096, "stop": None}

            num_chunks = math.ceil(len(text) / (4096 // 2))
            chunk_size = math.ceil(len(text) / num_chunks)

            chunks = []
            for i in range(0, len(text), chunk_size):
                chunks.append(text[i:i + chunk_size])

            results = []
            
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                query_data["messages"][0]["content"] = conversation_history_with_prompt + ' '.join(batch)
                results.extend(await asyncio.gather(*[send_chunk_to_gpt(url, headers, query_data) for _ in range(len(batch))]))

            response_data = results.json()
            #self.store_gpt_response_data(response_data)
            return response_data'''

    def provide_summary(self):
        if MainWindow.text==None:
            QMessageBox.warning(self, "Warning", "Please upload a document first!")
        else:
            progress_dialog = QProgressDialog("Summary getting created...", None, 0, 100 , self)
            progress_dialog.resize(400,100)
            progress_style = ProgressDialogStyle()
            progress_style.set_style(progress_dialog)
            progress_dialog.setWindowTitle("In progress")
            progress_dialog.setWindowModality(Qt.ApplicationModal)
            progress_dialog.setAutoClose(False)
            progress_dialog.setAutoReset(False)
            progress_dialog.setMinimumDuration(0)
            progress_dialog.forceShow()  

            for progress in range(0, 101, 3):
                progress_dialog.setValue(progress)
                QApplication.processEvents()
                progress_dialog.setLabelText(f"Please wait while the summary is being created...")
                QApplication.instance().processEvents()
                QThread.msleep(100)
            self.send_text_to_gpt(MainWindow.text)
            progress_dialog.close()
            if hasattr(self, "response_data"):
                summary_message = f"Response from Advisory GPT at {datetime.now().strftime('%m/%d/%Y %I:%M %p')}:\n\n"
                summary_text = self.response_data["choices"][0]["message"]["content"]
                print("summart text", summary_text)
                if '200' in str(MainWindow.summary_response):
                    MainWindow.gen_tst_case = False
                else:
                    summary_text = "Response not recieved. Please check GPT URL"

            
                MainWindow.gen_tst_case = False
                final_summary_text = f"{summary_message} {summary_text}"
                self.summary_label.setPlainText(final_summary_text)
                self.summary_label.setStyleSheet("color: black; background-color: white;")
                MainWindow.summary_button.setEnabled(False)
                MainWindow.summary_button.setStyleSheet("background-color: green; color: white; font-weight: bold; border-radius: 5px;")

                MainWindow.generate_test_cases_button.setEnabled(True)
                MainWindow.generate_test_cases_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")

                # MainWindow.generate_automation_scripts_button.setEnabled(True)
                # MainWindow.generate_automation_scripts_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")
        # else:
        #     QMessageBox.warning(self, "Warning", "Please upload a document first!")

    def generate_test_requirements(self):
        # MainWindow.generate_test_cases_button.setEnabled(True)
        # MainWindow.gen_tst_req = True
        if MainWindow.text==None:
            QMessageBox.warning(self, "Warning", "Please upload a document first!")
        else:    
        
            progress_dialog1 = QProgressDialog("Generating test requirements...", None, 0, 100, self)
            progress_dialog1.resize(400,100)
            progress_style = ProgressDialogStyle()
            progress_style.set_style(progress_dialog1)


            progress_dialog1.setWindowTitle("In progress")
            progress_dialog1.setWindowModality(Qt.ApplicationModal)
            progress_dialog1.setAutoClose(False)
            progress_dialog1.setAutoReset(False)
            progress_dialog1.setMinimumDuration(0)
            progress_dialog1.forceShow() 
    

            for progress in range(0, 101, 3):
                progress_dialog1.setValue(progress)
                QApplication.processEvents()
                progress_dialog1.setLabelText(f"Please wait while the test requirements are being generated...")
                QApplication.instance().processEvents()
                QThread.msleep(100)
            #self.send_text_to_gpt(MainWindow.text)
                
            self.api_key = "73c315ba0bc44e6abffd10b189dc3199"
            self.url = "https://gpt-llm-4.openai.azure.com/openai/deployments/gpt-4/chat/completions?api-version=2023-03-15-preview"  
            #self.api_key = "d4595b37d235460395fa78721602c87b"
            #self.url = "https://oai-kgsgpt-testapp.openai.azure.com/openai/deployments/gpt-35-turbo-16k/chat/completions?api-version=2023-03-15-preview"  
            '''
            conversation_history_with_prompt = f"""You are an expert in analyzing functional and non functional test case requirements. Understand the below text.
            -----
            {MainWindow.text}
            -----
            Generate all possible functional and non functional test Requirements in detail. 
            Requirements should be non-ambiguous and should be point wise.
            Each test case should be associated with a test case ID and should be in the below format:
            -----------
            <Test case ID>: <Test requirement 1>
            <test case ID>: <Test requirement 2>
            -----------
            A line should separate the functional and non functional requirements."""
            '''
            '''   
            conversation_history_with_prompt = f"""You are an expert in analyzing functional and non functional test case requirements. Understand the below text.
            -----
            {MainWindow.text}
            -----
            Generate all possible functional and non functional test Requirements in detail. 
            A line should separate the functional and non functional requirements.
            Each requirement should have a unique requirement ID and the requirements should be non-ambiguous and should be point wise.
            For each requirement, all possible sub requirements need to be generated point wise. 
            There can be one or more sub requirements under each requirement. 
            There should be no ambiguity or repitions.
            Each sub requirement should be associated with a unique test case ID.
    Do not generate any code.
    The output should be in plain text, point wise and prescise
            The output must be of the below format:
            -----------
            Requirement 1: <requirement ID> <Requirement description>
    <Test case ID>: <Sub-equirement 1 description>
    <test case ID>: <Sub-equirement 2 description>
    
            Requirement 2: <requirement ID> <Requirement description>
    <Test case ID>: <Sub-equirement 1 description>
    <test case ID>: <Sub-equirement 2 description>
            -----------
    The sub requirements descriptions must be related to the main requirement description
    There can be one or many sub requirements under each requirement.
    Requirements and sub requirements need to make sense and needs to be properly grouped."""
           '''
            conversation_history_with_prompt = f"""As a proficient analyst of functional and non-functional test case requirements, your task is to comprehend the following text and generate suitable requirements.
            -----
            {MainWindow.text}
            -----
            Prepare a comprehensive list of all possible functional and non-functional requirements in meticulous detail. Specific instructions:
            
            1. Maintain a distinct partition for functional and non-functional requirements.
            
            2. Each requirement should possess a unique ID, allowing for simpler identification and no possibility of duplicate entries. 

            3. Formulate the requirements in a manner that avoids any ambiguity- they should aim to be concise, precise and clearly stated.
            
            4. Within each requirement, you must generate a list of all possible sub-requirements. Assign a unique test case ID to each sub-requirement.

            5. Avoid generating any code.
            
            6. Your output should be formatted in plain text and must abode by the following pattern:

            -----------
            Requirement 1: <requirement ID> <Requirement description>
                <Test case ID>: <Sub-requirement 1 description>
                <Test case ID>: <Sub-requirement 2 description>

            Requirement 2: <requirement ID> <Requirement description>
                <Test case ID>: <Sub-requirement 1 description>
                <Test case ID>: <Sub-requirement 2 description>
            -----------
            
            7. Ensure the description of each sub-requirement is directly relevant to the main requirement.

            8. There may be multiple sub-requirements under each main requirement.
            
            9. Requirements and sub requirements must make sense and be properly categorized under each respective group.
            """

            headers = {"api-key": self.api_key}
            query_data = {
                "messages": [
                    {"role": "system", "content": conversation_history_with_prompt}
                ],
                "temperature": 0,
                "top_p": 1,
                "frequency_penalty": 0,
                "presence_penalty": 0,
                "max_tokens": 4096,
                "stop": None,
            }

            req_response = requests.post(self.url, headers=headers, json=query_data, verify=False)
            MainWindow.req_response_data = req_response.json()
            self.store_gpt_response_data(MainWindow.req_response_data)

            if hasattr(self, "response_data"):
                progress_dialog1.close()
                summary_message = f"Response from Advisory GPT at {datetime.now().strftime('%m/%d/%Y %I:%M %p')}:\n\n"
                choices = self.response_data.get("choices", [])
                '''
                if choices:
                        MainWindow.gen_tst_case = False
                        summary_text = choices[0]["message"]["content"]
                else:
                        summary_text = "Response not recieved. Please check GPT URL"
                '''
                if '200' in str(req_response):
                    MainWindow.gen_tst_case = False
                    summary_text = MainWindow.req_response_data["choices"][0]["message"]["content"]

                else:
                    summary_text = "Response not recieved. Please check GPT URL"
                MainWindow.req_txt = summary_text
                summary_text = f"{summary_message} {summary_text}"
                self.summary_label.setPlainText(summary_text)
                self.summary_label.setStyleSheet("color: black; background-color: white;")
                MainWindow.generate_test_requirements_button.setEnabled(False)
                MainWindow.generate_test_requirements_button.setStyleSheet("background-color: green; color: white; font-weight: bold; border-radius: 5px;")

                MainWindow.generate_test_cases_button.setEnabled(True)
                MainWindow.generate_test_cases_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")

                MainWindow.generate_automation_scripts_button.setEnabled(False)
                MainWindow.generate_automation_scripts_button.setStyleSheet("background-color: #B6B6B4; color: white; font-weight: bold; border-radius: 5px;")
                
                MainWindow.summary_button.setEnabled(False)
                MainWindow.summary_button.setStyleSheet("background-color: green; color: white; font-weight: bold; border-radius: 5px;")

        # else:
        #     QMessageBox.warning(self, "Warning", "Please upload a document first!")

    def generate_test_cases(self, prompt):

        # # For example:
        # for progress_val in range(0, 101, 3):
        #     QApplication.processEvents()
        #     progress = QProgressBar(self)
        #     if progress:
        #         progress.setValue(progress_val)
        #         QApplication.instance().processEvents()
        #         QThread.msleep(100)

        if hasattr(self, "response_data"):
            self.api_key = "73c315ba0bc44e6abffd10b189dc3199"
            self.url = "https://gpt-llm-4.openai.azure.com/openai/deployments/gpt-4/chat/completions?api-version=2023-03-15-preview"  
            #self.api_key = "d4595b37d235460395fa78721602c87b"
            #self.url = "https://oai-kgsgpt-testapp.openai.azure.com/openai/deployments/gpt-35-turbo-16k/chat/completions?api-version=2023-03-15-preview" 
            headers = {"api-key": self.api_key}
            query_data = {
                "messages": [
                    {"role": "system", "content": "You are a helpful assistant that generate testcases to validate different oracle applications"},
                    {"role": "user", "content": prompt}
                    
                ],
                "temperature": 0,
                "top_p": 1,
                "frequency_penalty": 0,
                "presence_penalty": 0,
                "max_tokens": 4096,
                "stop": None,
            }
            response = requests.post(self.url, headers=headers, json=query_data, verify=False)
            print("Generate test cases Response:",response)
            MainWindow.generate_test_response_data = response.json()
            #print("Response_data:",response_data)
            if '200' in str(response):
                MainWindow.gen_tst_case = True
                self.functional_test_cases = MainWindow.generate_test_response_data["choices"][0]["message"]["content"]
            else:
                test_cases_text = "Response not recieved. Please check GPT URL"

            MainWindow.test_cases = self.functional_test_cases
            test_cases_message = f"Test Cases generated by Advisory GPT at {datetime.now().strftime('%m/%d/%Y %I:%M %p')}:\n\n"
            test_cases_text = f"{test_cases_message} {self.functional_test_cases}"

            self.summary_label.setPlainText(test_cases_text)
            #self.summary_label.setPlainText(str(table))
            self.summary_label.setStyleSheet("color: black; background-color: white;")
            MainWindow.generate_test_cases_button.setEnabled(False)
            MainWindow.generate_test_cases_button.setStyleSheet("background-color: green; color: white; font-weight: bold; border-radius: 5px;")

            MainWindow.generate_automation_scripts_button.setEnabled(True)
            MainWindow.generate_automation_scripts_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")

            MainWindow.upload_to_opkey_button.setEnabled(True)
            MainWindow.upload_to_opkey_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")

            MainWindow.generate_test_requirements_button.setEnabled(False)
            MainWindow.generate_test_requirements_button.setStyleSheet("background-color: green; color: white; font-weight: bold; border-radius: 5px;")
        else:
            QMessageBox.warning(self, "Warning", "Please upload a document first!")

    def test_type(self):
        test_dialog = QDialog(self)
        test_dialog.resize(600, 150)
        dialog_style =ButtonDialogStyle()
        dialog_style.set_style(test_dialog)

        test_dialog.setWindowTitle("Select the type of test case")

        vbox = QVBoxLayout()
        test_dialog.setLayout(vbox)

        # label = QLabel("Select the type of test case:")
        # vbox.addWidget(label)

        combo = QComboBox()
        combo.setMinimumWidth(200)
        combo.setStyleSheet("background-color: white; border-radius: 5px; min-height: 1.3em; padding: 3px 10px")
        combo.move(60, 30)
        combo.addItem("(default) Select the type of testing...")
        combo.addItem("Functional Test Cases")
        combo.addItem("Non-Functional Test Cases")
        combo.addItem("Unit Test Cases")
        combo.addItem("Performance Test Cases")
        combo.addItem("Regression Test Cases")
        combo.addItem("Integration Test Cases")
        combo.addItem("Smoke Test Cases")
        combo.addItem("Security Test Cases")
        combo.addItem("End-to-end Test Cases")
        vbox.addWidget(combo)

        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        vbox.addWidget(button_box)

        button_box.accepted.connect(test_dialog.accept)
        button_box.rejected.connect(test_dialog.reject)

        result = test_dialog.exec_()

        if result == QDialog.Accepted:
            test_type = combo.currentText()
            MainWindow.generate_test_type_user_input = test_type
           
            progress_dialog2 = QProgressDialog("generating_test_cases...", None, 0, 100, self)
            progress_dialog2.resize(400,100)
            progress_style = ProgressDialogStyle()
            progress_style.set_style(progress_dialog2)

            progress_dialog2.setWindowTitle("In progress")
            progress_dialog2.setWindowModality(Qt.ApplicationModal)
            progress_dialog2.setAutoClose(False)
            progress_dialog2.setAutoReset(False)
            progress_dialog2.setMinimumDuration(0)
            progress_dialog2.forceShow()  

            for progress in range(0, 101, 3):
                    progress_dialog2.setValue(progress)
                    QApplication.processEvents()
                    progress_dialog2.setLabelText(f"Please wait while the test cases are generating..")
                    QApplication.instance().processEvents()
                    QThread.msleep(150)

            if MainWindow.gen_tst_req and (test_type=="Functional Test Cases" or test_type=="Non-Functional Test Cases"):
                
                prompt=''
                if test_type!="(default) Select the type of testing...":
                    #self.send_text_to_gpt(MainWindow.text)

            #   prompt=f"""You are an expert test engineer. Generate only the {test_type} in detail from the below text: 
            # ---------------
            # {MainWindow.req_txt} 
            # ---------------
            # The requirement ID in the response should should be the same as in the above text.
            # The requirement description in the response should should be the same as in the above text.
            # The test case ID in the response should should be the same as in the above text.
            # The number of test cases generated should be in the same as mentioned in the above text and should be in the same order as well.
            # Each test case should have: 'Requirement ID', 'Requirement Description',Test Case ID', 'Description', 'Detailed Steps', and 'Expected Results'.
            # The Detailed Steps should be as descriptive as possible and should be mentioned point wise.
            # Each test case should be separated.
            # Do not generate a code.
            # Sample test case:
            
            # 1. Requirement ID: <this should match with the requirement ID in the above text>
            #    Requirement Description: <this should match with the requirement description in the above text>
            #    Test Case ID: <this should match with the test case ID in the above text>
            #    Description: 
            #    Detailed Steps:
            #    Expected Result: 
            
            # All generated test cases should strictly adhere to above format.
            # """
        
                    prompt = f"""As a seasoned test engineer with a focus on SAP FIORI applications, your task is to generate {test_type} intuitively and meticulously. Use the following text as your source:
            ---------------
            {MainWindow.req_txt} 
            ---------------
            Balance consistency and relevance in your output. Ensure that the 'Requirement ID', 'Requirement Description' and 'Test Case ID' in your response match exactly with the information given in the text. 

            The test cases you generate should correlate with the numbers mentioned in the text, maintaining their order. Each case should provide information on 'Requirement ID', 'Requirement Description', 'Test Case ID', a clear 'Description', 'Detailed Steps', and 'Expected Results'. 
            
            Strive for descriptiveness and clarity in the 'Detailed Steps' section and split this information into points for easier comprehension. Don't forget to differentiate each test case clearly.

            Note that no additional code should be generated – we're solely focusing on the analysis and construction of test cases.
            Sample testcase:
            
Requirement ID: FR1
Requirement Description: Order Processing
Test Case ID: TC1
Description: Verify the functionality of sales order creation.
Detailed Steps:
 1.1. Launch the application.
 1.2. Enter the sales representative username and password in the respective fields.
 1.3. Click on the Login button.
 1.4. Verify successful login by checking for a welcome message or the user's name on the dashboard.
 2.1. On the user dashboard, locate and click on the Sales Order tab.
 2.2. From the dropdown menu, select "Create Sales Order".
 2.3. Confirm navigation to the sales order creation page by verifying the page title or header.
 3.1. On the sales order creation page, locate the order form.
 3.2. Enter valid data in all mandatory fields such as product, quantity, customer details, delivery details, etc.
 3.3. For non-mandatory fields, leave some empty to ensure they are not required for order creation.
 4.1. Click the Submit button to complete the order.
 4.2. Verify that an order confirmation message or order number is displayed on the screen.
 4.3. Check if the new order can be seen in the sales order list.
Expected Results: 

Successful login into the application.
Successful navigation to the sales order creation page.
Successful entry of all required details.
Confirmation of sales order creation is visible and the new sales order is listed in the sales order list.




            Your test case should adhere to the following sample format: 
            
            1. Requirement ID: <Matches with the Requirement ID in the text>
               Requirement Description: <Matches with the Requirement Description in the text>
               Test Case ID: <Matches with the Test Case ID in the text>
               Description: <A clear 'Description' that succinctly outlines what the test will achieve>
               Detailed Steps: <Detailed Steps' that chronologically explain how to execute the test, broken down point by point for easy comprehension and Screen-wise, executable steps that offer a simple explanation of how the test case is to be run>
               Expected Results: <'Expected Results' outlining what should happen upon successful completion of the test>
            
            Meticulous attention to the format is crucial – all produced test cases must follow this template precisely."""

                    self.generate_test_cases(prompt)
                    progress_dialog2.close()
                else:
                    prompt = None
                    QMessageBox.warningf
            else:
                prompt=''
                if test_type!="(default) Select the type of testing...":
            #         prompt=f"""You are an expert test engineer. Generate the maximum possible {test_type} in detail for: 
            # ---------------
            # {MainWindow.text} 
            # ---------------
            # Each test case should have: 'Requirement ID', 'Requirement Description','Test Case ID', 'Description', 'Detailed Steps' and 'Expected Results'.
            # The Detailed Steps should be as descriptive as possible and should be mentioned point wise.
            # Each test case should be separated.
            # Do not generate a code.
            # Sample test case:
            
            # 1. Requirement ID: 
            #    Requirement Description:
            #    Test Case ID: 
            #    Description: 
            #    Detailed Steps:
            #    Expected Result: 
            
            # All generated test cases should strictly adhere to above format.
            #"""
                    prompt = f"""As a seasoned test engineer with a focus on SAP FIORI applications, your task is to generate {test_type} intuitively and meticulously. Use the following text as your source:
            ---------------
            {MainWindow.req_txt} 
            ---------------
            Balance consistency and relevance in your output. Ensure that the 'Requirement ID', 'Requirement Description' and 'Test Case ID' in your response match exactly with the information given in the text. 

            The test cases you generate should correlate with the numbers mentioned in the text, maintaining their order. Each case should provide information on 'Requirement ID', 'Requirement Description', 'Test Case ID', a clear 'Description', 'Detailed Steps', and 'Expected Results'. 
            
            Strive for descriptiveness and clarity in the 'Detailed Steps' section and split this information into points for easier comprehension. Don't forget to differentiate each test case clearly.

            Note that no additional code should be generated – we're solely focusing on the analysis and construction of test cases.

            Your test case should adhere to the following sample format: 
            Sample testcase:
            
Requirement ID: FR1
Requirement Description: Order Processing
Test Case ID: TC1
Description: Verify the functionality of sales order creation.
Detailed Steps:
 1.1. Launch the application.
 1.2. Enter the sales representative username and password in the respective fields.
 1.3. Click on the Login button.
 1.4. Verify successful login by checking for a welcome message or the user's name on the dashboard.
 2.1. On the user dashboard, locate and click on the Sales Order tab.
 2.2. From the dropdown menu, select "Create Sales Order".
 2.3. Confirm navigation to the sales order creation page by verifying the page title or header.
 3.1. On the sales order creation page, locate the order form.
 3.2. Enter valid data in all mandatory fields such as product, quantity, customer details, delivery details, etc.
 3.3. For non-mandatory fields, leave some empty to ensure they are not required for order creation.
 4.1. Click the Submit button to complete the order.
 4.2. Verify that an order confirmation message or order number is displayed on the screen.
 4.3. Check if the new order can be seen in the sales order list.
Expected Results: 

Successful login into the application.
Successful navigation to the sales order creation page.
Successful entry of all required details.
Confirmation of sales order creation is visible and the new sales order is listed in the sales order list.




            1. Requirement ID: <Matches with the Requirement ID in the text>
               Requirement Description: <Matches with the Requirement Description in the text>
               Test Case ID: <Matches with the Test Case ID in the text>
               Description: <A clear 'Description' that succinctly outlines what the test will achieve>
               Detailed Steps: <Detailed Steps' that chronologically explain how to execute the test, broken down point by point for easy comprehension and Screen-wise, executable steps that offer a simple explanation of how the test case is to be run>
               Expected Results: <'Expected Results' outlining what should happen upon successful completion of the test>
            
            Meticulous attention to the format is crucial – all produced test cases must follow this template precisely."""

                    self.generate_test_cases(prompt)
                    progress_dialog2.close()
                else:
                    prompt = None
                    QMessageBox.warning(self, "Warning", "Please select the type of testing")
 
    def create_automation_scripts(self):
        lang_dialog = QDialog(self)
        lang_dialog.resize(600,150)
        lang_dialog_style =ButtonDialogStyle()
        lang_dialog_style.set_style(lang_dialog)
        lang_dialog.setWindowTitle("Select Automation Framework")

        vbox = QVBoxLayout()
        lang_dialog.setLayout(vbox)

        # label = QLabel("Select the automation framework:")
        # vbox.addWidget(label)

        combo_box = QComboBox()
        combo_box.setMinimumWidth(200)
        combo_box.setStyleSheet("background-color: white; border-radius: 5px; min-height: 1.3em; padding: 3px 10px")
        combo_box.move(60, 30)
        combo_box.addItem("(default) Select the automation framework...")
        combo_box.addItem("Selenium Java")
        combo_box.addItem("Selenium Python")
        combo_box.addItem("Python")
        combo_box.addItem("C#")
        combo_box.addItem("Julia")
        vbox.addWidget(combo_box)

        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        vbox.addWidget(button_box)

        button_box.accepted.connect(lang_dialog.accept)
        button_box.rejected.connect(lang_dialog.reject)

        result = lang_dialog.exec_()

        if result == QDialog.Accepted:
            MainWindow.automation_script = True
            framework = combo_box.currentText()
            progress_dialog3 = QProgressDialog("generate automation scripts...", None, 0, 100, self)
            progress_dialog3.resize(400,100)
            progress_style = ProgressDialogStyle()
            progress_style.set_style(progress_dialog3)
            
            progress_dialog3.setWindowTitle("In progress")
            progress_dialog3.setWindowModality(Qt.ApplicationModal)
            progress_dialog3.setAutoClose(False)
            progress_dialog3.setAutoReset(False)
            progress_dialog3.setMinimumDuration(0)
            progress_dialog3.forceShow()  

            for progress in range(0, 101, 3):
                progress_dialog3.setValue(progress)
                QApplication.processEvents()
                progress_dialog3.setLabelText(f"Please wait while the scripts are generating...")
                QApplication.instance().processEvents()
                QThread.msleep(100)

            if framework == "(default) Select the automation framework...":
                QMessageBox.warning(self, "Warning", "Please select an automation framework!")
            else:
                MainWindow.gen_tst_case = False
                self.generate_automation_scripts(framework)
                progress_dialog3.close()
        else:
            return


    def generate_automation_scripts(self, framework):
        if hasattr(self, "response_data") and hasattr(self, "functional_test_cases"):
                    self.api_key = "73c315ba0bc44e6abffd10b189dc3199"
                    self.url = "https://gpt-llm-4.openai.azure.com/openai/deployments/gpt-4/chat/completions?api-version=2023-03-15-preview"  

                    #self.api_key = "d4595b37d235460395fa78721602c87b"
                    #self.url = "https://oai-kgsgpt-testapp.openai.azure.com/openai/deployments/gpt-35-turbo-16k/chat/completions?api-version=2023-03-15-preview"  

                    conversation_history_with_prompt = f'''Convert the following functional test case into corresponding {framework} scripts, one for each test case:
                    -----
                    {MainWindow.test_cases}
                    -----
                    After converting the test case, provide detailed step by step instructions (point wise) on how to run the  automation script.
Separate the auomation scripts for each test case properly.
The output needs to be of the below format:
-----
Requirement ID:
Test case ID
Test case description:

Automation script
-----
Requirement ID, Test case ID and Test case description should match with above text.'''

                    headers = {"api-key": self.api_key}
                    query_data = {
                        "messages": [
                            {"role": "system", "content": conversation_history_with_prompt}
                        ],
                        "temperature": 0,
                        "top_p": 1,
                        "frequency_penalty": 0,
                        "presence_penalty": 0,
                        "max_tokens": 4096,
                        "stop": None,
                    }

                    response = requests.post(self.url, headers=headers, json=query_data, verify=False)
                    print("Automation script response:",response)
                    response_data = response.json()
                    if '200' in str(response):
                        MainWindow.gen_tst_case = False
                        automation_script_text = response_data["choices"][0]["message"]["content"]
                    else:
                        automation_script_text = "Response not recieved. Please check GPT URL. Could not generate automation scripts"
                    '''
                    choices2 = response_data.get("choices", [])
                    if choices2:
                        MainWindow.gen_tst_case = False
                        automation_script_text = choices2[0]["message"]["content"]
                    else:
                        automation_script_text = "Response not recieved. Please check GPT URL"
                    '''
                    #MainWindow.gen_tst_case = False
                    auto_message = f"Automation Scripts generated by Advisory GPT at {datetime.now().strftime('%m/%d/%Y %I:%M %p')}:\n\n"
                    automation_script_final = f"{auto_message} {automation_script_text}"
                    self.summary_label.setPlainText(automation_script_final)
                    self.summary_label.setStyleSheet("color: black;  background-color: white;")
                    MainWindow.generate_automation_scripts_button.setEnabled(False)
                    MainWindow.generate_automation_scripts_button.setStyleSheet("background-color: green; color: white; font-weight: bold; border-radius: 5px;")
        else:
            QMessageBox.warning(self, "Warning", "Please upload a document and generate functional test cases first!")

    def store_gpt_response_data(self, response_data):
        self.response_data = response_data

    def upload_to_opkey_function(self):
        progress_dialog = QProgressDialog("Summary getting created...", None, 0, 100 , self)
        progress_dialog.resize(400,100)
        progress_style = ProgressDialogStyle()
        progress_style.set_style(progress_dialog)
        progress_dialog.setWindowTitle("In progress")
        progress_dialog.setWindowModality(Qt.ApplicationModal)
        progress_dialog.setAutoClose(False)
        progress_dialog.setAutoReset(False)
        progress_dialog.setMinimumDuration(0)
        progress_dialog.forceShow()  

        for progress in range(0, 101, 3):
            progress_dialog.setValue(progress)
            QApplication.processEvents()
            progress_dialog.setLabelText(f"Please wait...")
            QApplication.instance().processEvents()
            QThread.msleep(100)
        
        login_cookies = self.login()
        print("Login response:", json.dumps(dict(login_cookies), indent=4))
    
        select_project_result = self.select_project(dict(login_cookies))
        print("Select Project response:", json.dumps(dict(select_project_result), indent=4))
    
        MainWindow.upload_to_opkey_text = self.create_test_case(dict(select_project_result))
        print("Create test case response:", json.dumps(MainWindow.upload_to_opkey_text, indent=4))
        MainWindow.upload_to_opkey_text = self.create_test_case(dict(select_project_result))
        print("Create test case response:", json.dumps(MainWindow.upload_to_opkey_text, indent=4))
        progress_dialog.close()
        # for _ in range(3):  # Retry 3 times
        # MainWindow.upload_to_opkey_text = self.create_test_case(dict(select_project_result))
        # print("Create test case response:", json.dumps(MainWindow.upload_to_opkey_text, indent=4))
        # if MainWindow.upload_to_opkey_text.strip() != "":
        #     break  # If the response is non-empty, break the loop
        # time.sleep(1)  # Wait for 1 second before the next retry

        self.summary_label.setPlainText(json.dumps(MainWindow.upload_to_opkey_text, indent=4))
        self.summary_label.setStyleSheet("color: black;  background-color: white;")
        MainWindow.upload_to_opkey_button.setEnabled(False)
        MainWindow.upload_to_opkey_button.setStyleSheet("background-color: green; color: white; font-weight: bold; border-radius: 5px;")


    def login(self):
        api_key = '5SOEH6MWEPX6NZ4IX0'
        headers = {'ApiKey': api_key}
        url = 'https://demo.labs.opkeyone.com/api/OpkeyAuth/Login_API_KEY?username=tuhinmitra@kpmg.com&loginsource=TCG'
    
        response = requests.post(url, headers=headers, verify=False)
        s=requests.Session()
        print(response.json())
    
        #print(response.cookies['ASP.NET_SessionId'])
        print(response.cookies)
    
        return response.cookies  # Return session ID
    
    def select_project(self, cookies):
        url = 'https://demo.labs.opkeyone.com/api/OpkeyAuth/SelectProject?projectId=9e188d9d-29eb-4dc5-af85-d39638f9e3df'
        response = requests.post(url,cookies=cookies,verify=False)
        return response.cookies    
    
    def create_test_case(self, cookies):
        url = 'https://demo.labs.opkeyone.com/Automation/ExecuteOracleTestCreationQuery'
        formdata = {
        "requirement": "Role :ACME Purchase Order Inquiry Role Description : Role consists of fundamental query functions and does not grant any permission to process transactions Task for the Role • View Purchase Order • View Item • Search Supplier",
        "manualTestCase": "Test case Description : Verify power user with \"ACME_PURCHASE_ORDER_INQUIRY_JOB\" role is able to view purchase order\n\nTest Steps :\n\n1) Login with a user which has access to \"ACME_PURCHASE_ORDER_INQUIRY_JOB\"\n\n2) Click on Procurement -> Click on Purchase Orders\n\n3) Click on the \"Tasks\" pane present on the right side of the screen.\n\n4) Select \"Manage Orders\" present under \"Orders\" in the task panel\n\n5) Clear all search fields\n\n6)Enter the Order number and Click on \"Search\"\n\n7)Click the Order number hyperlink to review the PO\n\n8)Click on \"Done\"\n\nExpected Result -\n\n1)User is able to login\n\n2) User is able to Navigate to Purchase order under Procurement\n\n3) User is able to click on the \"Tasks\"\n\n4) User is able to Navigate to the \"Manage Orders\" page and can see the Search input parameters.\n\n5) User is able to clear all the any values populated in the field.\n\n6) User is able to enter the Order Number in the Order field and Cick on Search button\n\n7) User should be able to view the Order details, and is able to click on the hyperlink\n\n8) User is able to view and click on Done button",
        }
        # Assuming multipart/form-data needed for this request
        response = requests.post(url, cookies=cookies,data=formdata, verify=False)
        return response.json()


    def export_button(self):

        if MainWindow.gen_tst_case:

            exp_dialog = QDialog(self)
            exp_dialog.resize(600,150)
            exp_dialog_style =ButtonDialogStyle()
            exp_dialog_style.set_style(exp_dialog)
            exp_dialog.setWindowTitle("Select download file format")

            vbox = QVBoxLayout()
            exp_dialog.setLayout(vbox)

            # label = QLabel("Select the download file format:")
            # vbox.addWidget(label)

            exp_combo_box = QComboBox()
            exp_combo_box.addItem("(default) Select the download file format...")
            exp_combo_box.addItem("Excel")
            exp_combo_box.addItem("Word")
            vbox.addWidget(exp_combo_box)

            button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
            vbox.addWidget(button_box)

            button_box.accepted.connect(exp_dialog.accept)
            button_box.rejected.connect(exp_dialog.reject)

            result = exp_dialog.exec_()
            if result == QDialog.Accepted:
                export_format = exp_combo_box.currentText()
                if export_format == "(default) Select the automation framework...":
                    QMessageBox.warning(self, "Warning", "Please select an automation framework!")

                input_string = MainWindow.generate_test_response_data["choices"][0]["message"]["content"]
                
                data = self.convert_testcases_into_dict(str(input_string))
                print("data is ",data)
                
                if export_format=='Word':
                    file_name, _ = QFileDialog.getSaveFileName(self, "Export to Word", "", "Word Documents (*.docx);;All Files (*)")

                    doc = docx.Document()
                    if file_name:

                        # Add a table to the document
                        table = doc.add_table(rows=len(data) + 1, cols=len(data[0]) + 4)  # Update the number of columns

                        # Format table
                        table.style = doc.styles['Table Grid']
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER

                        # Add header row
                        headers = ['Req ID', 'Req Description', 'Test Case ID', 'Description', 'Detailed Steps', 'Expected Result',
                                'Delivery Team', 'Type', 'Priority', 'Milestone']  # Update headers list
                        for idx, header in enumerate(headers):
                            table.cell(0, idx).text = header
                            table.cell(0, idx).paragraphs[0].runs[0].font.bold = True
                            table.cell(0, idx).paragraphs[0].runs[0].font.size = Pt(12)
                            table.cell(0, idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # Populate table with data
                        for row_num, entry in enumerate(data, start=1):
                            print("row_num:", row_num)
                            print("entry", entry)
                            for col_num, key in enumerate(entry):
                                table.cell(int(row_num), int(col_num)).text = str(entry[key])
                                table.cell(int(row_num), int(col_num)).paragraphs[0].runs[0].font.size = Pt(12)
                                table.cell(int(row_num), int(col_num)).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                            # Write user input values in the current row
                            table.cell(row_num, 6).text = MainWindow.delivery_team_user_input
                            table.cell(row_num, 7).text = MainWindow.generate_test_type_user_input
                            table.cell(row_num, 8).text = MainWindow.priority_user_input
                            table.cell(row_num, 9).text = MainWindow.milestone_user_input

                        # Set column widths
                        for col in table.columns:
                            col.width = docx.shared.Inches(1.5)

                        # Save document
                        # progress_dialog4.close()
                        doc.save(file_name)
                    
                        QMessageBox.information(self, "Success", "Content successfully exported to Word!")

            
                elif export_format == 'Excel':
                    # create a new Excel workbook with a single worksheet
                    workbook = openpyxl.Workbook()
                    worksheet = workbook.active

                    # Write headers to first row of worksheet
                    headers = ['Req ID', 'Req Description', 'Test Case ID', 'Description', 'Detailed Steps', 'Expected Result',
                            'Delivery Team', 'Type', 'Priority', 'Milestone']
                    for col_num, header in enumerate(headers, start=1):
                        worksheet.cell(row=1, column=col_num).value = header
                        worksheet.cell(row=1, column=col_num).font = openpyxl.styles.Font(bold=True, size=12)

                    # Adjust column widths
                    worksheet.column_dimensions['E'].width = 50
                    worksheet.column_dimensions['B'].width = 50
                    worksheet.column_dimensions['D'].width = 50
                    worksheet.column_dimensions['F'].width = 25
                    worksheet.column_dimensions['G'].width = 25

                    # Write each row of data to a new row in the worksheet
                    for row_num, row_data in enumerate(data, start=2):
                        for col_num, key in enumerate(row_data, start=1):
                            if key == "Detailed Steps":
                                cell = worksheet.cell(row=row_num, column=col_num)
                                cell.value = row_data[key]
                                cell.alignment = openpyxl.styles.Alignment(wrap_text=True)
                            else:
                                worksheet.cell(row=row_num, column=col_num).value = row_data[key]
                                worksheet.cell(row=row_num, column=col_num).font = openpyxl.styles.Font(size=12)

                        # Write user input values in the current row
                        worksheet.cell(row=row_num, column=7).value = MainWindow.delivery_team_user_input
                        worksheet.cell(row=row_num, column=8).value = MainWindow.generate_test_type_user_input
                        worksheet.cell(row=row_num, column=9).value = MainWindow.priority_user_input
                        worksheet.cell(row=row_num, column=10).value = MainWindow.milestone_user_input

                    # progress_dialog4.close()
                    # Save the workbook
                    file_name, _ = QFileDialog.getSaveFileName(self, "Export to Excel", "", "Excel Documents (*.xlsx);;All Files (*)")
                    if file_name:
                        workbook.save(file_name + ".xlsx")
                        
                        QMessageBox.information(self, "Success", "Content successfully exported to Excel!")
                    else:
                        file_name, _ = QFileDialog.getSaveFileName(self, "Export to Excel", "", "Excel Documents (*.xlsx);;All Files (*)")

            
            
                elif export_format == 'Excel':
                    # create a new Excel workbook with a single worksheet
                    workbook = openpyxl.Workbook()
                    worksheet = workbook.active

                    # Write headers to first row of worksheet
                    headers = ['Req ID', 'Req Description', 'Test Case ID', 'Description', 'Detailed Steps', 'Expected Result',
                            'Delivery Team', 'Type', 'Priority', 'Milestone']
                    for col_num, header in enumerate(headers, start=1):
                        worksheet.cell(row=1, column=col_num).value = header
                        worksheet.cell(row=1, column=col_num).font = openpyxl.styles.Font(bold=True, size=12)

                    # Adjust column widths
                    worksheet.column_dimensions['E'].width = 50
                    worksheet.column_dimensions['B'].width = 50
                    worksheet.column_dimensions['D'].width = 50
                    worksheet.column_dimensions['F'].width = 25
                    worksheet.column_dimensions['G'].width = 25

                    # Write each row of data to a new row in the worksheet
                    for row_num, row_data in enumerate(data, start=2):
                        for col_num, key in enumerate(row_data, start=1):
                            if key == "Detailed Steps":
                                cell = worksheet.cell(row=row_num, column=col_num)
                                cell.value = row_data[key]
                                cell.alignment = openpyxl.styles.Alignment(wrap_text=True)
                            else:
                                worksheet.cell(row=row_num, column=col_num).value = row_data[key]
                                worksheet.cell(row=row_num, column=col_num).font = openpyxl.styles.Font(size=12)

                        # Write user input values in the current row
                        worksheet.cell(row=row_num, column=7).value = MainWindow.delivery_team_user_input
                        worksheet.cell(row=row_num, column=8).value = MainWindow.generate_test_type_user_input
                        worksheet.cell(row=row_num, column=9).value = MainWindow.priority_user_input
                        worksheet.cell(row=row_num, column=10).value = MainWindow.milestone_user_input

                    # progress_dialog4.close()
                    # Save the workbook
                    file_name, _ = QFileDialog.getSaveFileName(self, "Export to Excel", "", "Excel Documents (*.xlsx);;All Files (*)")
                    if file_name:
                        workbook.save(file_name + ".xlsx")
                        
                        QMessageBox.information(self, "Success", "Content successfully exported to Excel!")
                    else:
                        file_name, _ = QFileDialog.getSaveFileName(self, "Export to Excel", "", "Excel Documents (*.xlsx);;All Files (*)")

   
                else:
                    return
        else:
                file_name, _ = QFileDialog.getSaveFileName(self, "Export to Word", "", "Word Documents (*.docx);;All Files (*)")
                if file_name:    
                    doc = Document()
                    doc.add_paragraph(self.summary_label.toPlainText())
                    doc.save(file_name)
                    QMessageBox.information(self, "Success", "Content successfully exported to Word!")
        
      
        MainWindow.export_file_button.setEnabled(True)
        MainWindow.export_file_button.setStyleSheet("background-color: green; color: white; font-weight: bold; border-radius: 5px;")
        
    def export_to_word(self):
        if MainWindow.gen_tst_case:
            input_string = MainWindow.generate_test_response_data["choices"][0]["message"]["content"]
            #self.api_key = "d4595b37d235460395fa78721602c87b"
            #self.url = "https://oai-kgsgpt-testapp.openai.azure.com/openai/deployments/gpt-35-turbo-16k/chat/completions?api-version=2023-03-15-preview"  
            progress_dialog4 = QProgressDialog("Downloading document...", None, 0, 100, self)
            progress_dialog4.resize(400,100)
            progress_style = ProgressDialogStyle()
            progress_style.set_style(progress_dialog4)
            progress_dialog4.setWindowTitle("In progress")
            progress_dialog4.setWindowModality(Qt.ApplicationModal)
            progress_dialog4.setAutoClose(False)
            progress_dialog4.setAutoReset(False)
            progress_dialog4.setMinimumDuration(0)
            progress_dialog4.forceShow()  

            for progress in range(0, 101, 3):
                progress_dialog4.setValue(progress)
                QApplication.processEvents()
                progress_dialog4.setLabelText(f"Please wait while the document is downloading...")
                QApplication.instance().processEvents()
                QThread.msleep(150)
            

            self.api_key = "73c315ba0bc44e6abffd10b189dc3199"
            self.url = "https://gpt-llm-4.openai.azure.com/openai/deployments/gpt-4/chat/completions?api-version=2023-03-15-preview"  

            conversation_history_with_prompt = f"""I have data in below format
        -----
        {input_string}
        -----
        Output this string in the form a single python list of dictionaries.
Each dictionary in the list should represent a test case and should have below keys:
'Req ID', 'Req Description', 'Test Case ID', 'Description', 'Detailed Steps', 'Expected Result'.
Req ID is the same as Requirement ID.
Req Description is the same as Requirement Description.
Case of the keys should not matter, they can be in any case, but should match keys mentioned.
In case any value has new lines, the new lines should be retained.
I do not want any code, I just want to see list of dictionaries created."""

            headers = {"api-key": self.api_key}
            query_data = {
                "messages": [
                    {"role": "system", "content": conversation_history_with_prompt}
                ],
                "temperature": 0,
                "top_p": 1,
                "frequency_penalty": 0,
                "presence_penalty": 0,
                "max_tokens": 4096,
                "stop": None,
            }

            req_response = requests.post(self.url, headers=headers, json=query_data, verify=False)
            list_of_dict_response_data = req_response.json()
            test_cases = list_of_dict_response_data["choices"][0]["message"]["content"]

            # Removing anything before '['
            open_bracket_position = test_cases.find('[')
            rem_line_1 = test_cases[open_bracket_position:]

            # Removing anything after last ']'
            close_bracket_position=rem_line_1.rfind(']')
            rem_line_last = rem_line_1[:close_bracket_position + 1]
            print("rem_line_last: ",rem_line_last)
            #rem_line_last = rem_line_last.replace("\'", "\"") # Replace single quotes with double quotes
            #print("rem_line_last final: ",rem_line_1)
            list_of_dicts = ast.literal_eval(rem_line_last)

            data = []
# Print the list of dictionaries
            for d in list_of_dicts:
                data.append(d)
            
            print("data is ",data)

            file_name, _ = QFileDialog.getSaveFileName(self, "Export to Word", "", "Word Documents (*.docx);;All Files (*)")

            doc = docx.Document()
            if file_name:

    # Add a table to the document
                table = doc.add_table(rows=len(data) + 1, cols=len(data[0]))

    # Format table
                table.style = doc.styles['Table Grid']
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add header row
                headers = ['Req ID','Req Description','Test Case ID', 'Description', 'Detailed Steps', 'Expected Result']
                for idx, header in enumerate(headers):
                    table.cell(0, idx).text = header
                    table.cell(0, idx).paragraphs[0].runs[0].font.bold = True
                    table.cell(0, idx).paragraphs[0].runs[0].font.size = Pt(12)
                    table.cell(0, idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Populate table with data
                for row_num, entry in enumerate(data, start=1):
                    print("row_num:",row_num)
                    print("entry",entry)
                    for col_num, key in enumerate(entry):
                        table.cell(int(row_num), int(col_num)).text = str(entry[key])
                        table.cell(int(row_num), int(col_num)).paragraphs[0].runs[0].font.size = Pt(12)
                        table.cell(int(row_num), int(col_num)).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set column widths
                for col in table.columns:
                    col.width = docx.shared.Inches(1.5)

    # Save document
                doc.save(file_name)
                progress_dialog4.close()
                QMessageBox.information(self, "Success", "Content successfully exported to Word!")
            progress_dialog4.close()
        else:
            file_name, _ = QFileDialog.getSaveFileName(self, "Export to Word", "", "Word Documents (*.docx);;All Files (*)")
            if file_name:
                doc = Document()
                doc.add_paragraph(self.summary_label.toPlainText())
                doc.save(file_name)
                QMessageBox.information(self, "Success", "Content successfully exported to Word!")

    def clear_display(self):
        self.summary_label.setPlainText("")
        MainWindow.upload_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")
        MainWindow.summary_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")
        MainWindow.generate_test_requirements_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")
        # MainWindow.generate_test_cases_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")
        # MainWindow.generate_automation_scripts_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")
        MainWindow.generate_test_cases_button.setStyleSheet("background-color: #B6B6B4; color: white; font-weight: bold; border-radius: 5px;")
        MainWindow.generate_automation_scripts_button.setStyleSheet("background-color: #B6B6B4; color: white; font-weight: bold; border-radius: 5px;")
        MainWindow.user_input_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")
        MainWindow.export_file_button.setStyleSheet("background-color: #29465B; color: white; font-weight: bold; border-radius: 5px;")
             
 
        MainWindow.upload_button.setEnabled(True)
        MainWindow.summary_button.setEnabled(True)
        MainWindow.generate_test_requirements_button.setEnabled(True)
        # MainWindow.generate_test_cases_button.setEnabled(True)
        MainWindow.generate_test_cases_button.setEnabled(False)
        # MainWindow.generate_automation_scripts_button.setEnabled(True)
        MainWindow.generate_automation_scripts_button.setEnabled(False)
        MainWindow.user_input_button.setEnabled(True)
        MainWindow.upload_to_opkey_button(False)

        MainWindow.text = None
        MainWindow.generate_test_type_user_input = "..."
        MainWindow.delivery_team_user_input = "..."
        MainWindow.priority_user_input = "..."
        MainWindow.milestone_user_input ='...'

        MainWindow.req_response_data = None
        MainWindow.generate_test_response_data = None
        MainWindow.gen_tst_case = None
        MainWindow.gen_tst_req = None
        MainWindow.summary_response = None
        MainWindow.req_txt = None
        MainWindow.test_cases = None
        MainWindow.upload_to_opkey_text = None
               
    def exit_application(self):
        self.close()

    def download_templates(self):
        template_dialog = QDialog(self)
        template_dialog.resize(500,100)
        template_dialog_style =ButtonDialogStyle()
        template_dialog_style.set_style(template_dialog)
        template_dialog.setWindowTitle("Select The Template")

        vbox = QVBoxLayout()
        template_dialog.setLayout(vbox)

        # label = QLabel("Select the template that you want to download:")
        # vbox.addWidget(label)

        temp_combo_box = QComboBox()
        temp_combo_box.addItem("(default) Select the template...")
        temp_combo_box.addItem("Requirement template")
        temp_combo_box.addItem("User story template")
        temp_combo_box.addItem("Sample generic document")
        vbox.addWidget(temp_combo_box)

        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        vbox.addWidget(button_box)

        button_box.accepted.connect(template_dialog.accept)
        button_box.rejected.connect(template_dialog.reject)

        result = template_dialog.exec_()

        if result == QDialog.Accepted:
            selected_template = temp_combo_box.currentText()
            template_content=""
            if selected_template == "(default) Select the template...":
                QMessageBox.warning(self, "Warning", "Please select an automation framework!")
            elif selected_template == "Requirement template":
                file_name, _ = QFileDialog.getSaveFileName(self, "Export to Word", "", "Word Documents (*.docx);;All Files (*)")
                if file_name:
                    template_content="""
ABC application calculates the factorial of a number.
Below are the functional and non-functional requirements of the application.
-----------
Functional Requirements:

Requirement 1: FR1 - The program should correctly calculate the factorial of a number.
Test case ID: TC1.1 - Verify if the program correctly calculates the factorial of a positive integer.
Test case ID: TC1.2 - Verify if the program returns 1 when the input is 0, as the factorial of 0 is 1.
Test case ID: TC1.3 - Verify if the program handles negative numbers correctly, as factorials of negative numbers are undefined.

Requirement 2: FR2 - The program should correctly handle non-numeric inputs.
Test case ID: TC2.1 - Verify if the program handles string inputs correctly by displaying an appropriate error message.
Test case ID: TC2.2 - Verify if the program handles special character inputs correctly by displaying an appropriate error message.

Requirement 3: FR3 - The program should correctly handle large number inputs.
Test case ID: TC3.1 - Verify if the program correctly calculates the factorial of a large number without any errors or delays.
Test case ID: TC3.2 - Verify if the program handles numbers larger than the maximum limit correctly by displaying an appropriate error message.

-----------
Non-Functional Requirements:

Requirement 1: NFR1 - The program should have a user-friendly interface.
Test case ID: NTC1.1 - Verify if the program prompts the user to enter a number in a clear and understandable manner.
Test case ID: NTC1.2 - Verify if the program displays the result in a clear and understandable manner.

Requirement 2: NFR2 - The program should have a quick response time.
Test case ID: NTC2.1 - Verify if the program calculates the factorial of a number and displays the result within an acceptable time frame.

Requirement 3: NFR3 - The program should handle errors gracefully.
Test case ID: NTC3.1 - Verify if the program displays a clear and understandable error message when an error occurs.
Test case ID: NTC3.2 - Verify if the program does not crash or freeze when an error occurs.

Note: Requirement ID, description, Test Case ID and Test case description is mandatory.

"""
                    doc = Document()
                    doc.add_paragraph(template_content)
                    doc.save(file_name)
                    QMessageBox.information(self, "Success", "Downloaded 'Requirement template' successfully!")
            
            elif selected_template == "User story template":
                file_name, _ = QFileDialog.getSaveFileName(self, "Export to Word", "", "Word Documents (*.docx);;All Files (*)")
                if file_name:
                    template_content="""
Title: Factorial Calculator for Students

User Story (JIRA format):
As a high school student studying mathematics, I want a simple and user-friendly program to quickly calculate the factorials of given numbers, So that I can save time and verify my answers without manual calculations.

Acceptance Criteria:
1.	The program must prompt the user to enter a number.
2.	The program must be able to handle positive integer inputs, including 0.
3.	The program must calculate the factorial of the entered number correctly.
4.	The program must display the calculated factorial in a clear and understandable format.
5.	Errors or exceptions should be handled gracefully with appropriate error messages for invalid inputs, like negative numbers or non-integer values.

Story Label: mathematics, factorial 

Story Type: Story 

Story Points: 3 

Story Priority: Medium 

Due date (if applicable): MM/DD/YYYY 

Assignee: John Doe
"""
                    doc = Document()
                    doc.add_paragraph(template_content)
                    doc.save(file_name)
                    QMessageBox.information(self, "Success", "Downloaded 'User story template' successfully!")
            
            else:
                file_name, _ = QFileDialog.getSaveFileName(self, "Export to Word", "", "Word Documents (*.docx);;All Files (*)")
                if file_name:
                    template_content="""
This document has a program to generate the factorial of a number.
-----------------------------------------------------------------
def factorial(num):
    fact = 1
    for i in range(1, num + 1):
        fact = fact * i
    return fact

num = int(input("Enter a number: "))
print("Factorial of", num, "is", factorial(num))
-----------------------------------------------------------------
"""
                    doc = Document()
                    doc.add_paragraph(template_content)
                    doc.save(file_name)
                    QMessageBox.information(self, "Success", "Downloaded 'Sample generic document' successfully!")

    def user_input(self):
        doc_dialog = QDialog(self)
        doc_dialog.resize(500, 200)
        doc_dialog_style =ButtonDialogStyle()
        doc_dialog_style.set_style(doc_dialog)
        
        doc_dialog.setWindowTitle("User Inputs")
 
        # Add grid layout to the document dialog box
        grid = QGridLayout()
        doc_dialog.setLayout(grid)
        # Add widgets to the grid
        label1 = QLabel("Delivery Team")
        label1.setStyleSheet("QLabel { color: white; font-weight: bold; }")
        grid.addWidget(label1, 0, 0)
 
        field1 = QComboBox()
        field1.addItems(["(default) Select the Delivery Team","FI - OTC","FI - OTC RAR","FI - PTP","FI - RTR"
                         ,"FI - RTR Tax","FI - TAX","FI - ITR","FI - Financial Close","Blank","Input"])
        
        grid.addWidget(field1, 0, 1)
        

        label2 = QLabel("Type")
        label2.setStyleSheet("QLabel { color: white; font-weight: bold; }")
        grid.addWidget(label2, 1, 0)
 
        # field2 = QComboBox()
        # field2.addItems(["(default) Select the type","Functional Test Cases","Non-Functional Test Cases","Unit Test Cases",
        # "Performance Test Cases","Regression Test Cases","Integration Test Cases","Smoke Test Cases",
        # "Security Test Cases","End-to-end Test Cases"])
        field2 = QLabel(MainWindow.generate_test_type_user_input)
        field2.setStyleSheet("QLabel { color: white;  font-size: 16px; font-weight: bold; qproperty-alignment: 'AlignLeft';  }")
        grid.addWidget(field2, 1, 1)
        # # Create a QFrame to use as the background box
        # box = QFrame()
        # box.setFixedSize(350,28)
        # # box.setStyleSheet("QFrame { background-color: white;}")

        # # Create a layout for the box
        # box_layout = QVBoxLayout()
        # box_layout.addWidget(field2)
        # box.setLayout(box_layout)

        # # Add the box to the grid
        # grid.addWidget(box, 1, 1)
 
        label3 = QLabel("Priority:")
        label3.setStyleSheet("QLabel { color: white; font-weight: bold; }")
        grid.addWidget(label3, 2, 0)
 
        field3 = QComboBox()
        field3.addItems(["Medium", "High", "Low"])
        grid.addWidget(field3, 2, 1)
 
        label4 = QLabel("Milestone:")
        label4.setStyleSheet("QLabel { color: white; font-weight: bold; }")
        grid.addWidget(label4, 3, 0)
 
        field4 = QLineEdit()
        grid.addWidget(field4, 3, 1)
 
 
        # # Add Ok and Cancel buttons to the dialog box
        # btn_ok = QPushButton("Ok")
        # btn_cancel = QPushButton("Cancel")
        # grid.addWidget(btn_ok, 4, 0)
        # grid.addWidget(btn_cancel, 4, 1)

        # Create an horizontal layout for Ok and Cancel buttons
        button_layout = QHBoxLayout()
 
        btn_ok = QPushButton("Ok")
        btn_cancel = QPushButton("Cancel")
 
        # Set buttons to have the same size
        btn_ok.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        btn_cancel.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
 
        button_layout.addStretch(1)  # Add a stretchable space before the buttons
        button_layout.addWidget(btn_ok)
        button_layout.addWidget(btn_cancel)
 
        grid.addLayout(button_layout, 4, 1)  # Add the horizontal layout to the grid layout at position (4, 1)
 
        # Connect the buttons to their respective actions
        btn_ok.clicked.connect(doc_dialog.accept)
        btn_cancel.clicked.connect(doc_dialog.reject)
 
        # Connect the buttons to their respective actions
        btn_ok.clicked.connect(doc_dialog.accept)
        btn_cancel.clicked.connect(doc_dialog.reject)
 
        if doc_dialog.exec_() == QDialog.Accepted:
            # Process the input data
            MainWindow.delivery_team_user_input= field1.currentText()
            type= MainWindow.generate_test_type_user_input
            MainWindow.priority_user_input = field3.currentText()
            MainWindow.milestone_user_input = field4.text()

            if MainWindow.delivery_team_user_input == "(default) Select the Delivery Team" and MainWindow.milestone_user_input =="" and type =="..." and MainWindow.priority_user_input =='Medium':
                MainWindow.delivery_team_user_input= "..."
                MainWindow.priority_user_input = "..."
                MainWindow.milestone_user_input = "..."
            if MainWindow.delivery_team_user_input == "(default) Select the Delivery Team":
                MainWindow.delivery_team_user_input= "..."
 
            print("Delivery team:", MainWindow.delivery_team_user_input)
            print("Type:",type)
            print("Priority:", MainWindow.priority_user_input)
            print("Milestone:", MainWindow.milestone_user_input)

            MainWindow.user_input_button.setEnabled(False)
            MainWindow.user_input_button.setStyleSheet("background-color: green; color: white; font-weight: bold; border-radius: 5px;")

    # def convert_testcases_into_dict(self,text):
    #     data = text.split("\n")
    #     test_cases = []

    #     req_id = None
    #     req_desc = None
    #     steps = None  
    #     for line in data:
    #         line = line.strip()
    #         if ':' not in line: 
    #             continue

    #         if 'requirement id' in line.lower() or 'req id' in line.lower():
    #             req_id = line.split(':')[1].strip()
    #         elif 'requirement description' in line.lower() or 'req description' in line.lower():
    #             req_desc = line.split(':')[1].strip()
    #         elif 'test case id' in line.lower() or 'case id' in line.lower():
    #             test_id = line.split(':')[1].strip()
    #         elif 'description' in line.lower():
    #             desc = line.split(':')[1].strip()
    #         elif 'detailed steps' in line.lower():

    #             #steps = [line]
    #             steps = [line.split(':')[1].strip()][1:]
    #             # steps = [line.split(':')[1].strip()] if ':' in line else [line]

         

    #         elif 'expected result' in line.lower():
    #             result = line.split(':')[1].strip()
    #             steps.append('')  # add an empty line after the last step to preserve new lines
    #             print(steps)
    #             test_cases.append({
    #                 'Req ID': req_id,
    #                 'Req Description': req_desc,
    #                 'Test Case ID': test_id,
    #                 'Description': desc,
    #                 'Detailed Steps': '\n'.join(steps).strip(),
    #                 'Expected Result': result
    #             })
    #         elif steps is not None:
    #             steps.append(line)
    #         # elif steps is not None and ':' in line:
    #         #     sep_line = line.split(':')[1].strip() if ':' in line else line
    #         #     steps.append(sep_line.split()[-1] if sep_line[0].isnumeric() else sep_line)


    #     # Print the list of dictionaries
    #     print(test_cases)
    #     return test_cases
            
    def convert_testcases_into_dict(self, text):
        data = text.split("\n")
        test_cases = []

        req_id = None
        req_desc = None
        test_id = None
        desc = None
        steps = None
        for line in data:
            line = line.strip()
            if ':' not in line and steps is None:
                continue

            if 'requirement id' in line.lower() or 'req id' in line.lower():
                req_id = line.split(':')[1].strip()
            elif 'requirement description' in line.lower() or 'req description' in line.lower():
                req_desc = line.split(':')[1].strip()
            elif 'test case id' in line.lower() or 'case id' in line.lower():
                test_id = line.split(':')[1].strip()
            elif 'description' in line.lower():
                desc = line.split(':')[1].strip()
            elif 'detailed steps' in line.lower():
                # steps = [line.split(':', 1)[1].strip()]
                steps = [line.split(':')[1].strip()] if ':' in line else [line]
            elif 'expected result' in line.lower():
                result = line.split(':')[1].strip()
                steps.append('')  # add an empty line after the last step to preserve new lines
                test_cases.append({
                    'Req ID': req_id,
                    'Req Description': req_desc,
                    'Test Case ID': test_id,
                    'Description': desc,
                    'Detailed Steps': '\n'.join(steps).strip(),
                    'Expected Result': result
                })
                steps = None
            elif steps is not None:
                steps.append(line)

        return test_cases


if __name__ == '__main__':
    app = QApplication(sys.argv)

    main_window = MainWindow()

    if main_window.show_popup():
        main_window.show()
        sys.exit(app.exec_())
    else:
        sys.exit(0)